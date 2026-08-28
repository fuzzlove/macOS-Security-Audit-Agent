from __future__ import annotations

import csv
import html
import json
import textwrap
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable

from mac_audit_agent.keylogger_detection import KeyloggerScanReport


REPORT_FORMATS: tuple[tuple[str, str, str], ...] = (
    ("txt", "Plain Text — Primary Static Handoff", "Text (*.txt)"),
    ("pdf", "PDF — Primary Read-Only Brief", "PDF (*.pdf)"),
    ("csv", "CSV — Sanitized Findings Data", "CSV (*.csv)"),
    ("json", "JSON — Structured Evidence Data", "JSON (*.json)"),
    ("html", "HTML — Static Review Copy", "HTML (*.html)"),
    ("docx", "DOCX — Macro-Free Department Brief", "Word Document (*.docx)"),
    ("xlsx", "XLSX — Formula-Free Contractor Workbook", "Excel Workbook (*.xlsx)"),
)

DISCLAIMER = (
    "This report documents observed keyboard-event-tap and privacy-permission indicators. "
    "A capability or permission is not, by itself, proof of keylogging, malware, attribution, or compromise. "
    "MSAA does not collect keystrokes. Findings require authorized human validation and correlation."
)


def build_keylogger_report(report: KeyloggerScanReport) -> dict[str, Any]:
    findings = [finding.to_dict() for finding in report.findings]
    severity = Counter(str(item.get("severity", "unknown")).lower() for item in findings)
    confidence = Counter(str(item.get("confidence", "unknown")).lower() for item in findings)
    classification = Counter(str(item.get("classification", "unknown")) for item in findings)
    scores = [int(item.get("score", 0) or 0) for item in findings]
    confidence_percentages = [int(item.get("analytic_confidence_percent", 0) or 0) for item in findings]
    false_positive_percentages = [int(item.get("false_positive_risk_percent", 100) or 0) for item in findings]
    return {
        "report_title": "Keylogger Detection Threat Assessment",
        "report_purpose": "Authorized incident-response, security-operations, management, legal, and contractor review.",
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "scan_period": {"started_at": report.started_at, "completed_at": report.completed_at},
        "executive_summary": {
            "finding_count": len(findings),
            "high_or_critical_count": severity["high"] + severity["critical"],
            "event_tap_count": report.event_tap_count,
            "relevant_privacy_grant_count": report.tcc_grant_count,
            "maximum_threat_score": max(scores, default=0),
            "average_threat_score": round(sum(scores) / len(scores), 2) if scores else 0,
            "average_analytic_confidence_percent": round(sum(confidence_percentages) / len(confidence_percentages), 2) if confidence_percentages else 0,
            "average_false_positive_risk_percent": round(sum(false_positive_percentages) / len(false_positive_percentages), 2) if false_positive_percentages else 0,
            "measured_accuracy_rate_percent": report.accuracy_rate_percent,
            "accuracy_basis": report.accuracy_basis,
        },
        "threat_statistics": {
            "by_severity": dict(sorted(severity.items())),
            "by_confidence": dict(sorted(confidence.items())),
            "by_classification": dict(sorted(classification.items())),
        },
        "coverage": dict(report.coverage),
        "coverage_warnings": list(report.warnings),
        "findings": findings,
        "threat_knowledge": report.threat_knowledge,
        "handling_notice": DISCLAIMER,
        "report_safety": "Report-only output: no macros, scripts, formulas, embedded executables, or captured keystrokes are included.",
    }


def _safe_cell(value: Any) -> str:
    text = value if isinstance(value, str) else json.dumps(value, sort_keys=True, default=str)
    return "'" + text if text.lstrip().startswith(("=", "+", "-", "@")) else text


def _lines(payload: dict[str, Any]) -> list[str]:
    summary = payload["executive_summary"]
    lines = [payload["report_title"], "=" * len(payload["report_title"]), "", payload["report_purpose"], "", "HANDLING AND ANALYTIC QUALIFICATION", payload["handling_notice"], payload["report_safety"], "", "EXECUTIVE SUMMARY"]
    lines.extend(f"{key.replace('_', ' ').title()}: {value}" for key, value in summary.items())
    lines.extend(["", "SCAN PERIOD", f"Started: {payload['scan_period']['started_at']}", f"Completed: {payload['scan_period']['completed_at']}", f"Report generated: {payload['generated_at']}", "", "THREAT STATISTICS"])
    for category, values in payload["threat_statistics"].items():
        lines.append(f"{category.replace('_', ' ').title()}: " + ", ".join(f"{key}={value}" for key, value in values.items()))
    lines.extend(["", "COLLECTION COVERAGE"])
    lines.extend(f"{key}: {value}" for key, value in payload["coverage"].items())
    if payload["coverage_warnings"]:
        lines.extend(["", "COVERAGE WARNINGS", *[f"- {item}" for item in payload["coverage_warnings"]]])
    lines.extend(["", "DETAILED FINDINGS"])
    if not payload["findings"]:
        lines.append("No findings were reported by this scan.")
    for index, finding in enumerate(payload["findings"], 1):
        lines.extend([
            "", f"Finding {index}: {finding.get('title', '')}", "-" * 72,
            f"Finding ID: {finding.get('finding_id', '')}", f"Severity: {str(finding.get('severity', '')).upper()}",
            f"Confidence: {finding.get('confidence', '')}", f"Threat score: {finding.get('score', 0)}/100",
            f"Analytic confidence: {finding.get('analytic_confidence_percent', 0)}%",
            f"Estimated false-positive risk: {finding.get('false_positive_risk_percent', 100)}%",
            f"Percentage basis: {finding.get('percentage_basis', '')}",
            f"Classification: {finding.get('classification', '')}", f"Process/client: {finding.get('process_name') or finding.get('bundle_id') or 'Unresolved'}",
            f"PID: {finding.get('pid') or 'Not available'}", f"Path: {finding.get('path') or 'Not resolved'}",
            "Signals:", *[f"- {signal}" for signal in finding.get("signals", [])],
            f"Recommended action: {finding.get('recommendation', '')}",
            "Intervention: " + "; ".join(finding.get("intervention_actions", [])),
            "Removal: " + "; ".join(finding.get("removal_actions", [])),
            "Remediation: " + "; ".join(finding.get("remediation_actions", [])),
            "MITRE ATT&CK context: " + json.dumps(finding.get("attack_techniques", []), sort_keys=True, default=str),
            "Documented threat context (not attribution): " + json.dumps(finding.get("documented_threat_context", []), sort_keys=True, default=str),
            "Evidence: " + json.dumps(finding.get("evidence", {}), sort_keys=True, default=str),
        ])
    return lines


def export_keylogger_report(report: KeyloggerScanReport, destination: Path, format_id: str) -> Path:
    destination = Path(destination)
    payload = build_keylogger_report(report)
    exporters: dict[str, Callable[[dict[str, Any], Path], None]] = {
        "txt": _export_txt, "pdf": _export_pdf, "csv": _export_csv, "json": _export_json,
        "html": _export_html, "docx": _export_docx, "xlsx": _export_xlsx,
    }
    if format_id not in exporters:
        raise ValueError(f"Unsupported Keylogger Detection report format: {format_id}")
    destination.parent.mkdir(parents=True, exist_ok=True)
    exporters[format_id](payload, destination)
    return destination


def _export_txt(payload: dict[str, Any], path: Path) -> None:
    path.write_text("\n".join(_lines(payload)) + "\n", encoding="utf-8")


def _export_json(payload: dict[str, Any], path: Path) -> None:
    path.write_text(json.dumps(payload, indent=2, sort_keys=True, default=str) + "\n", encoding="utf-8")


def _export_csv(payload: dict[str, Any], path: Path) -> None:
    fields = ["finding_id", "severity", "confidence", "score", "analytic_confidence_percent", "false_positive_risk_percent", "classification", "process_or_client", "pid", "path", "signals", "recommendation", "intervention", "removal", "remediation", "attack_techniques", "evidence"]
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields)
        writer.writeheader()
        for finding in payload["findings"]:
            row = {
                "finding_id": finding.get("finding_id", ""), "severity": finding.get("severity", ""), "confidence": finding.get("confidence", ""),
                "score": finding.get("score", 0), "classification": finding.get("classification", ""),
                "analytic_confidence_percent": finding.get("analytic_confidence_percent", 0), "false_positive_risk_percent": finding.get("false_positive_risk_percent", 100),
                "process_or_client": finding.get("process_name") or finding.get("bundle_id") or "", "pid": finding.get("pid") or "", "path": finding.get("path", ""),
                "signals": "; ".join(finding.get("signals", [])), "recommendation": finding.get("recommendation", ""),
                "intervention": "; ".join(finding.get("intervention_actions", [])), "removal": "; ".join(finding.get("removal_actions", [])), "remediation": "; ".join(finding.get("remediation_actions", [])),
                "attack_techniques": finding.get("attack_techniques", []), "evidence": finding.get("evidence", {}),
            }
            writer.writerow({key: _safe_cell(value) for key, value in row.items()})


def _export_html(payload: dict[str, Any], path: Path) -> None:
    content = "\n".join(f"<p>{html.escape(line)}</p>" if line else "<br>" for line in _lines(payload))
    path.write_text(f"<!doctype html><html><head><meta charset='utf-8'><title>{html.escape(payload['report_title'])}</title><style>body{{font:14px -apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;max-width:1000px;margin:40px auto;color:#17202a}}p{{margin:5px 0;white-space:pre-wrap}}h1{{color:#183b56}}</style></head><body><h1>{html.escape(payload['report_title'])}</h1>{content}</body></html>", encoding="utf-8")


def _export_pdf(payload: dict[str, Any], path: Path) -> None:
    # A deliberately minimal static PDF: built-in font, text only, and no actions,
    # attachments, JavaScript, forms, multimedia, or external converter dependency.
    wrapped: list[str] = []
    for line in _lines(payload):
        wrapped.extend(textwrap.wrap(line, width=100, replace_whitespace=False, drop_whitespace=False) or [""])
    pages = [wrapped[index:index + 55] for index in range(0, len(wrapped), 55)] or [[""]]
    page_count = len(pages)
    font_object = 3 + page_count * 2
    objects: dict[int, bytes] = {
        1: b"<< /Type /Catalog /Pages 2 0 R >>",
        font_object: b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    }
    page_refs = []
    for index, lines in enumerate(pages):
        page_object = 3 + index * 2; content_object = page_object + 1
        page_refs.append(f"{page_object} 0 R")
        commands = ["BT", "/F1 9 Tf", "42 750 Td", "12 TL"]
        for line in lines:
            safe = line.encode("latin-1", "replace").decode("latin-1").replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
            commands.append(f"({safe}) Tj T*")
        commands.append("ET")
        stream = "\n".join(commands).encode("latin-1")
        objects[page_object] = f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 {font_object} 0 R >> >> /Contents {content_object} 0 R >>".encode("ascii")
        objects[content_object] = f"<< /Length {len(stream)} >>\nstream\n".encode("ascii") + stream + b"\nendstream"
    objects[2] = f"<< /Type /Pages /Kids [{' '.join(page_refs)}] /Count {page_count} >>".encode("ascii")
    document = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"); offsets = [0]
    for object_id in range(1, font_object + 1):
        offsets.append(len(document)); document.extend(f"{object_id} 0 obj\n".encode("ascii") + objects[object_id] + b"\nendobj\n")
    xref = len(document); document.extend(f"xref\n0 {font_object + 1}\n0000000000 65535 f \n".encode("ascii"))
    for offset in offsets[1:]: document.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    document.extend(f"trailer\n<< /Size {font_object + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode("ascii"))
    path.write_bytes(document)


def _export_docx(payload: dict[str, Any], path: Path) -> None:
    try:
        from docx import Document
    except (ImportError, OSError):
        from mac_audit_agent.professional_report import structured_payload_report
        structured_payload_report(path, title=payload["report_title"], payload=payload, qualification=payload["handling_notice"])
        return
    document = Document(); document.core_properties.title = payload["report_title"]; document.core_properties.author = "macOS Security Audit Agent"
    document.add_heading(payload["report_title"], 0)
    for line in _lines(payload)[3:]:
        if line in {"EXECUTIVE SUMMARY", "SCAN PERIOD", "THREAT STATISTICS", "COLLECTION COVERAGE", "COVERAGE WARNINGS", "DETAILED FINDINGS", "HANDLING AND ANALYTIC QUALIFICATION"}:
            document.add_heading(line.title(), level=1)
        elif line.startswith("Finding "):
            document.add_heading(line, level=2)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line:
            document.add_paragraph(line)
    document.save(path)


def _export_xlsx(payload: dict[str, Any], path: Path) -> None:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill
    except (ImportError, OSError):
        from mac_audit_agent.professional_report import structured_payload_report
        structured_payload_report(path, title=payload["report_title"], payload=payload, qualification=payload["handling_notice"])
        return
    workbook = Workbook(); summary = workbook.active; summary.title = "Executive Summary"
    summary.append([payload["report_title"]]); summary["A1"].font = Font(size=16, bold=True)
    summary.append(["Handling notice", payload["handling_notice"]]); summary.append(["Report safety", payload["report_safety"]])
    for key, value in payload["executive_summary"].items(): summary.append([key.replace("_", " ").title(), _safe_cell(value)])
    for key, value in payload["coverage"].items(): summary.append([f"Coverage: {key}", _safe_cell(value)])
    findings = workbook.create_sheet("Findings")
    headers = ["Finding ID", "Severity", "Confidence", "Score", "Analytic Confidence %", "False-Positive Risk %", "Classification", "Process / Client", "PID", "Path", "Signals", "Recommendation", "Intervention", "Removal", "Remediation", "ATT&CK Context", "Evidence"]
    findings.append(headers)
    for cell in findings[1]: cell.font = Font(bold=True, color="FFFFFF"); cell.fill = PatternFill("solid", fgColor="183B56")
    for item in payload["findings"]:
        values = [item.get("finding_id", ""), item.get("severity", ""), item.get("confidence", ""), item.get("score", 0), item.get("analytic_confidence_percent", 0), item.get("false_positive_risk_percent", 100), item.get("classification", ""), item.get("process_name") or item.get("bundle_id") or "", item.get("pid") or "", item.get("path", ""), "; ".join(item.get("signals", [])), item.get("recommendation", ""), "; ".join(item.get("intervention_actions", [])), "; ".join(item.get("removal_actions", [])), "; ".join(item.get("remediation_actions", [])), item.get("attack_techniques", []), item.get("evidence", {})]
        findings.append([_safe_cell(value) for value in values])
    findings.freeze_panes = "A2"; findings.auto_filter.ref = findings.dimensions
    for sheet in workbook.worksheets:
        for column in sheet.columns:
            letter = column[0].column_letter; sheet.column_dimensions[letter].width = min(max(len(str(cell.value or "")) for cell in column) + 2, 60)
    workbook.save(path)

from __future__ import annotations

from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Any

from mac_audit_agent.assessment import SecurityAssessment
from mac_audit_agent.exporters.export_models import ExportAssessmentData, ExportOptions, build_export_assessment_data
from mac_audit_agent.ui.severity_styles import display_severity_label, get_severity_style


DISCLAIMER = (
    "MSAA provides security assessment and framework mapping support for analyst review. "
    "This report does not constitute certification, authorization, compliance approval, or a formal government assessment."
)


def _require_docx():
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("Word export requires python-docx. Install project dependencies including python-docx to create .docx reports.") from exc
    return Document, WD_ALIGN_PARAGRAPH, Pt, RGBColor


def _severity_color(severity: str):
    _Document, _align, _Pt, RGBColor = _require_docx()
    color = get_severity_style(severity).foreground.lstrip("#")
    return RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))


def _shade_risk_cell(cell, label: str) -> None:
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        return
    colors = get_severity_style(label)
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), colors.background.lstrip("#"))
    tc_pr.append(shading)


def _add_heading(doc, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_key_value_table(doc, rows: list[tuple[str, Any]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Value"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = str(value if value not in {None, ""} else "Not collected")


def _add_findings_table(doc, findings: list[dict[str, Any]], limit: int | None = None) -> None:
    rows = findings[:limit] if limit else findings
    if not rows:
        doc.add_paragraph("No findings were recorded for this section.")
        return
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for index, header in enumerate(["ID", "Severity", "Title", "Evidence Summary", "Suggested Fix"]):
        table.rows[0].cells[index].text = header
    for finding in rows:
        cells = table.add_row().cells
        cells[0].text = str(finding.get("finding_id", ""))
        severity = str(finding.get("severity", "info") or "info")
        _shade_risk_cell(cells[1], severity)
        sev = cells[1].paragraphs[0].add_run(display_severity_label(severity))
        sev.bold = True
        sev.font.color.rgb = _severity_color(severity)
        cells[2].text = str(finding.get("title", ""))
        cells[3].text = str(finding.get("evidence_summary") or finding.get("description", ""))
        cells[4].text = str(finding.get("suggested_fix", "Review."))


def _add_detailed_findings(doc, findings: list[dict[str, Any]]) -> None:
    if not findings:
        doc.add_paragraph("No findings were recorded for this section.")
        return
    fields = [
        ("ID", "finding_id"),
        ("Severity", "severity"),
        ("Confidence", "confidence"),
        ("Category", "category"),
        ("Description", "description"),
        ("Evidence", "evidence_summary"),
        ("Impact", "impact"),
        ("Suggested Fix", "suggested_fix"),
        ("Validation Steps", "validation_step"),
        ("False Positive Notes", "false_positive_notes"),
        ("NIST CSF", "nist_csf"),
        ("NIST 800-53", "nist_800_53"),
        ("MITRE ATT&CK", "mitre_attack"),
        ("CVE / CISA KEV", "cve"),
        ("Status", "status"),
    ]
    for finding in findings:
        _add_heading(doc, str(finding.get("title") or finding.get("finding_id") or "Finding"), 2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Field"
        table.rows[0].cells[1].text = "Value"
        for label, key in fields:
            value = finding.get(key, "")
            if key == "cve":
                cisa = finding.get("cisa_kev", "")
                value = ", ".join(str(item) for item in [value, cisa] if item)
            cells = table.add_row().cells
            cells[0].text = label
            if key == "severity":
                severity = str(value or "info")
                _shade_risk_cell(cells[1], severity)
                run = cells[1].paragraphs[0].add_run(display_severity_label(severity))
                run.bold = True
                run.font.color.rgb = _severity_color(severity)
            else:
                cells[1].text = str(value if value not in {None, ""} else "Not collected")


def _add_remediation(doc, data: ExportAssessmentData) -> None:
    _add_heading(doc, "Suggested Remediation Plan", 1)
    for priority in ["Immediate", "Short-Term", "Routine"]:
        _add_heading(doc, priority, 2)
        items = [item for item in data.remediation_items if item.get("priority") == priority]
        if not items:
            doc.add_paragraph("No items in this priority group.")
            continue
        for item in items:
            p = doc.add_paragraph(style=None)
            p.add_run(f"{item.get('finding_id', '')}: ").bold = True
            p.add_run(str(item.get("recommended_fix", "")))
            doc.add_paragraph(f"Validation: {item.get('validation_step', '')}")
            doc.add_paragraph(f"Difficulty: {item.get('difficulty', '')} | Expected impact: {item.get('expected_impact', '')}")


def _add_summary_section(doc, title: str, rows: list[dict[str, Any]]) -> None:
    _add_heading(doc, title, 1)
    if not rows:
        doc.add_paragraph("No data was collected for this section.")
        return
    for item in rows:
        text = item.get("summary") or item.get("evidence") or item.get("status") or item.get("name") or "Recorded item"
        doc.add_paragraph(str(text))
        if item.get("update_guidance_summary"):
            doc.add_paragraph(f"Update guidance: {item.get('update_guidance_summary')}")
        if item.get("verification_steps"):
            doc.add_paragraph(f"Verification steps: {item.get('verification_steps')}")
        if item.get("official_references"):
            doc.add_paragraph(f"Official references: {item.get('official_references')}")


def _build_document(data: ExportAssessmentData):
    Document, WD_ALIGN_PARAGRAPH, Pt, _RGBColor = _require_docx()
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MSAA Security Assessment")
    run.bold = True
    run.font.size = Pt(24)
    doc.add_paragraph("Generated by macOS Security Audit Agent").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Liquidsky Network Security").alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_key_value_table(
        doc,
        [
            ("Hostname", data.metadata.get("hostname", "")),
            ("macOS Version", data.metadata.get("macos_version", "")),
            ("Assessment Date", data.metadata.get("assessment_date", "")),
            ("Assessment ID", data.metadata.get("assessment_id", "")),
            ("Confidentiality", data.metadata.get("confidentiality_notice", "")),
        ],
    )
    doc.add_page_break()

    _add_heading(doc, "Executive Summary", 1)
    doc.add_paragraph(str(data.summary.get("executive_summary", "")))
    _add_key_value_table(
        doc,
        [
            ("Overall Risk Level", data.summary.get("risk_level", "")),
            ("Overall Score", data.summary.get("overall_score", "Unavailable")),
            ("Critical Findings", data.summary.get("critical_count", 0)),
            ("High Findings", data.summary.get("high_count", 0)),
            ("Medium Findings", data.summary.get("medium_count", 0)),
            ("Low Findings", data.summary.get("low_count", 0)),
            ("Info Findings", data.summary.get("info_count", 0)),
            ("Monitor Status", data.summary.get("monitor_status", "")),
            ("Apple Exposure Status", data.summary.get("apple_exposure_status", "")),
        ],
    )

    _add_heading(doc, "Risk Summary", 1)
    _add_key_value_table(
        doc,
        [
            ("Critical Count", data.summary.get("critical_count", 0)),
            ("High Count", data.summary.get("high_count", 0)),
            ("Medium Count", data.summary.get("medium_count", 0)),
            ("Low Count", data.summary.get("low_count", 0)),
            ("Info Count", data.summary.get("info_count", 0)),
        ],
    )

    _add_heading(doc, "Assessment Scope", 1)
    _add_key_value_table(
        doc,
        [
            ("Hostname", data.metadata.get("hostname", "")),
            ("macOS Version", data.metadata.get("macos_version", "")),
            ("Generated By", data.metadata.get("generated_by", "")),
            ("Data Sources", "Latest assessment, scan findings, monitor events, framework mappings, and subsystem summaries where available."),
            ("Limitations", "; ".join(data.limitations) if data.limitations else "None recorded."),
        ],
    )

    _add_heading(doc, "Top Findings", 1)
    _add_findings_table(doc, data.findings, limit=10)

    _add_heading(doc, "Detailed Findings", 1)
    _add_detailed_findings(doc, data.findings)

    _add_remediation(doc, data)
    _add_summary_section(doc, "Apple Exposure Assessment", data.apple_exposure)
    _add_summary_section(doc, "Network Activity Summary", data.network_activity)
    _add_summary_section(doc, "Admin and Persistence Summary", data.admin_persistence)
    _add_summary_section(doc, "Physical Device Summary", data.physical_devices)
    _add_summary_section(doc, "Visibility Integrity Summary", data.visibility_integrity)
    _add_summary_section(doc, "Application Integrity Verification", data.application_integrity)

    _add_heading(doc, "Framework Mapping Summary", 1)
    doc.add_paragraph("Mapped to, aligned with, and supports review of frameworks for analyst context. This wording does not imply certification, compliance, government approval, or authorization.")
    if data.framework_mappings:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        for idx, header in enumerate(["Finding ID", "Framework", "Control / Technique", "Name", "Notes"]):
            table.rows[0].cells[idx].text = header
        for item in data.framework_mappings[:100]:
            cells = table.add_row().cells
            cells[0].text = str(item.get("finding_id", ""))
            cells[1].text = str(item.get("framework", ""))
            cells[2].text = str(item.get("control_id", ""))
            cells[3].text = str(item.get("name", ""))
            cells[4].text = str(item.get("notes", ""))
    else:
        doc.add_paragraph("No framework mappings were available.")

    _add_heading(doc, "Appendix", 1)
    doc.add_paragraph("Limitations")
    for item in data.limitations or ["None recorded."]:
        doc.add_paragraph(str(item), style="List Bullet")
    doc.add_paragraph(DISCLAIMER)
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.text = "Generated by macOS Security Audit Agent"
    return doc


def export_assessment_word(assessment: SecurityAssessment, output_path: Path | None = None, *, options: ExportOptions | None = None) -> Path:
    data = build_export_assessment_data(assessment, options=options)
    if output_path is None:
        host = "".join(ch if ch.isalnum() or ch in {"-", "_"} else "_" for ch in assessment.hostname) or "mac"
        stamp = assessment.created_at.replace(":", "").replace("+", "Z").replace("-", "")[:15]
        from mac_audit_agent.reporting import get_reports_dir

        path = get_reports_dir() / f"MSAA_Security_Assessment_{host}_{stamp}.docx"
    else:
        path = output_path
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = _build_document(data)
    with NamedTemporaryFile(prefix=path.stem, suffix=".tmp.docx", dir=path.parent, delete=False) as handle:
        temp_path = Path(handle.name)
    try:
        doc.save(temp_path)
        temp_path.replace(path)
    except Exception:
        temp_path.unlink(missing_ok=True)
        raise
    return path

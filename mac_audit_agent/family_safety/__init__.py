from __future__ import annotations

import plistlib
import subprocess
import html
import json
from dataclasses import asdict, dataclass, field
from datetime import datetime
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Any

from mac_audit_agent.family_safety.categories import (
    FamilySafetyCategory,
    canonical_family_safety_categories,
    category_score,
    lockdown_plus_status,
)


SAFETY_CATEGORIES = [category.title for category in canonical_family_safety_categories()]

STATUSES = {"configured", "partially_configured", "not_configured"}


@dataclass
class FamilySafetyFinding:
    category: str
    title: str
    status: str
    summary: str
    recommendation: str
    evidence: str = ""

    def to_dict(self) -> dict[str, str]:
        return asdict(self)


@dataclass
class FamilySafetyScore:
    score: int
    categories: dict[str, int]
    recommended_improvements: list[str] = field(default_factory=list)
    completed_actions: list[str] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass
class SafetyEducationCard:
    topic: str
    guidance: str
    action: str

    def to_dict(self) -> dict[str, str]:
        return asdict(self)


@dataclass
class FamilySafetyReport:
    generated_at: str
    score: FamilySafetyScore
    findings: list[FamilySafetyFinding]
    wizard_recommendations: dict[str, list[str]]
    accessibility_checklist: list[FamilySafetyFinding]
    parent_checklist: list[FamilySafetyFinding]
    safe_browsing_status: list[FamilySafetyFinding]
    education_cards: list[SafetyEducationCard]
    app_review: list[FamilySafetyFinding]
    caregiver_dashboard: dict[str, Any]
    family_security_forecast: list[SafetyEducationCard]
    privacy_notice: list[str]
    category_definitions: list[FamilySafetyCategory] = field(default_factory=list)
    category_scores: dict[str, int] = field(default_factory=dict)
    government_lockdown_score: int = 0
    lockdown_plus_score: int = 0
    lockdown_plus_status: str = "Review Required"
    user_notes: dict[str, str] = field(default_factory=dict)
    completed_actions: dict[str, list[str]] = field(default_factory=dict)
    limitations: list[str] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return {
            "generated_at": self.generated_at,
            "score": self.score.to_dict(),
            "findings": [item.to_dict() for item in self.findings],
            "wizard_recommendations": self.wizard_recommendations,
            "accessibility_checklist": [item.to_dict() for item in self.accessibility_checklist],
            "parent_checklist": [item.to_dict() for item in self.parent_checklist],
            "safe_browsing_status": [item.to_dict() for item in self.safe_browsing_status],
            "education_cards": [item.to_dict() for item in self.education_cards],
            "app_review": [item.to_dict() for item in self.app_review],
            "caregiver_dashboard": self.caregiver_dashboard,
            "family_security_forecast": [item.to_dict() for item in self.family_security_forecast],
            "privacy_notice": self.privacy_notice,
            "category_definitions": [item.to_dict() for item in self.category_definitions],
            "category_scores": self.category_scores,
            "government_lockdown_score": self.government_lockdown_score,
            "lockdown_plus_score": self.lockdown_plus_score,
            "lockdown_plus_status": self.lockdown_plus_status,
            "user_notes": self.user_notes,
            "completed_actions": self.completed_actions,
            "limitations": self.limitations,
        }


class FamilySafetyAuditor:
    def __init__(self, home: Path | None = None) -> None:
        self.home = home or Path.home()

    def build_report(self, profile: str = "Shared Family Computer") -> FamilySafetyReport:
        findings = self._child_safety_findings()
        accessibility = self._accessibility_checklist()
        parent_checklist = self._parent_checklist(findings)
        safe_browsing = self._safe_browsing_status()
        app_review = self._app_review()
        all_scored = findings + accessibility + safe_browsing + app_review
        score = self._score(all_scored)
        category_definitions = self._category_definitions(all_scored)
        category_scores = {category.category_id: category_score(category, self._statuses_for_category(category, all_scored)) for category in category_definitions}
        government_score = category_scores.get("government_nist_lockdown", 0)
        lockdown_score = category_scores.get("lockdown_mode_plus", 0)
        return FamilySafetyReport(
            generated_at=datetime.now().isoformat(timespec="seconds"),
            score=score,
            findings=findings,
            wizard_recommendations=self.recommendations_for_profile(profile),
            accessibility_checklist=accessibility,
            parent_checklist=parent_checklist,
            safe_browsing_status=safe_browsing,
            education_cards=self.education_cards(),
            app_review=app_review,
            caregiver_dashboard=self._caregiver_dashboard(score, findings, app_review),
            family_security_forecast=self.family_security_forecast(),
            privacy_notice=self.privacy_notice(),
            category_definitions=category_definitions,
            category_scores=category_scores,
            government_lockdown_score=government_score,
            lockdown_plus_score=lockdown_score,
            lockdown_plus_status=lockdown_plus_status(lockdown_score),
            completed_actions=self._completed_actions_by_category(all_scored),
            limitations=[
                "Family Control Center is defensive, local-first guidance. It does not inspect private messages, browser history, screenshots, keystrokes, camera feeds, microphone audio, or private browsing data.",
                "NIST entries are mapped hardening guidance for review support only. This report does not claim compliance, certification, authorization, or approval for government use.",
                "Lockdown Mode Plus does not replace Apple Lockdown Mode and cannot guarantee protection.",
            ],
        )

    def _category_definitions(self, findings: list[FamilySafetyFinding]) -> list[FamilySafetyCategory]:
        categories = canonical_family_safety_categories()
        status_by_title: dict[str, str] = {}
        last_reviewed = datetime.now().isoformat(timespec="seconds")
        for finding in findings:
            status_by_title.setdefault(finding.category, finding.status)
        aliases = {
            "Screen Time": "Screen Time and Usage Controls",
            "Content Restrictions": "Screen Time and Usage Controls",
            "Privacy": "Privacy Permissions",
            "Downloads": "Downloads and File Safety",
            "Messaging": "Communication and Messaging Safety",
            "Web Safety": "Web and Browser Safety",
            "Location Sharing": "Privacy Permissions",
            "App Permissions": "Privacy Permissions",
            "Application Review": "App Store and Application Controls",
            "Accessibility": "Special Needs and Accessibility Safety",
            "System Security": "Government / NIST Lockdown Profile",
        }
        for category in categories:
            statuses = self._statuses_for_category(category, findings, aliases)
            if statuses:
                category.current_status = self._rollup_status(statuses)
                category.last_reviewed_at = last_reviewed
            elif category.category_id in {"government_nist_lockdown", "lockdown_mode_plus"}:
                category.current_status = "manual_verification_required"
                category.last_reviewed_at = last_reviewed
        return categories

    def _statuses_for_category(
        self,
        category: FamilySafetyCategory,
        findings: list[FamilySafetyFinding],
        aliases: dict[str, str] | None = None,
    ) -> list[str]:
        aliases = aliases or {
            "Screen Time": "Screen Time and Usage Controls",
            "Content Restrictions": "Screen Time and Usage Controls",
            "Privacy": "Privacy Permissions",
            "Downloads": "Downloads and File Safety",
            "Messaging": "Communication and Messaging Safety",
            "Web Safety": "Web and Browser Safety",
            "Location Sharing": "Privacy Permissions",
            "App Permissions": "Privacy Permissions",
            "Application Review": "App Store and Application Controls",
            "Accessibility": "Special Needs and Accessibility Safety",
            "System Security": "Government / NIST Lockdown Profile",
        }
        statuses = []
        for finding in findings:
            mapped_title = aliases.get(finding.category, finding.category)
            if mapped_title == category.title:
                statuses.append(finding.status)
        if category.category_id == "lockdown_mode_plus":
            statuses.append("manual_verification_required")
        return statuses

    def _rollup_status(self, statuses: list[str]) -> str:
        if any(status == "not_configured" for status in statuses):
            return "needs_review"
        if any(status in {"partially_configured", "manual_verification_required"} for status in statuses):
            return "manual_verification_required"
        return "configured"

    def _completed_actions_by_category(self, findings: list[FamilySafetyFinding]) -> dict[str, list[str]]:
        completed: dict[str, list[str]] = {}
        for finding in findings:
            if finding.status == "configured":
                completed.setdefault(finding.category, []).append(finding.title)
        return completed

    def _child_safety_findings(self) -> list[FamilySafetyFinding]:
        return [
            self._exists_finding("Screen Time", "Screen Time enabled", self.home / "Library" / "Application Support" / "com.apple.ScreenTimeAgent", "Open System Settings > Screen Time and confirm limits are enabled for the user."),
            self._verify("Account Safety", "Family Sharing configured", "Confirm Family Sharing is configured in Apple Account settings for child accounts.", "Family Sharing membership is not inspected by this tool."),
            self._verify("Content Restrictions", "Content & Privacy Restrictions enabled", "Open Screen Time > Content & Privacy and turn on restrictions appropriate for the user.", "Screen Time restrictions may be stored in protected system databases."),
            self._verify("Downloads", "App Store restrictions", "Use Screen Time to limit app installation, deletion, and in-app purchases when needed.", "This tool does not make App Store account changes."),
            self._verify("Content Restrictions", "Explicit content restrictions", "Review music, movies, apps, books, Siri, and search content ratings in Screen Time.", "Content settings require user confirmation."),
            self._verify("Web Safety", "Safari restrictions", "Use Screen Time web content controls and Safari privacy settings.", "Browsing history and website contents are never inspected."),
            self._verify("Web Safety", "Adult website filtering", "Turn on Limit Adult Websites or an approved allow-list for younger users.", "Website visits are not collected."),
            self._verify("Downloads", "App installation controls", "Require approval or admin credentials before new apps can be installed.", "Installation policy depends on account and Screen Time settings."),
            self._verify("Account Safety", "Password requirements", "Use a non-empty login password and avoid sharing administrator credentials.", "Password contents are never inspected."),
            self._guest_account_status(),
            self._verify("Location Sharing", "Location sharing settings", "Review Find My and Location Services with the user present.", "Location data is not read or stored."),
            self._verify("App Permissions", "Camera permissions", "Open Privacy & Security > Camera and remove apps that do not need access.", "Camera feeds are never accessed."),
            self._verify("App Permissions", "Microphone permissions", "Open Privacy & Security > Microphone and remove apps that do not need access.", "Microphone audio is never accessed."),
            self._command_status("System Security", "FileVault status", ["/usr/bin/fdesetup", "status"], "FileVault is On.", "Turn on FileVault to protect data if the Mac is lost."),
            self._command_status("System Security", "Firewall status", ["/usr/libexec/ApplicationFirewall/socketfilterfw", "--getglobalstate"], "enabled", "Turn on the firewall in Network settings."),
            self._sharing_status("System Security", "Remote login status", "Remote Login: On", "Keep Remote Login off unless a trusted admin needs it."),
            self._sharing_status("System Security", "Screen sharing status", "Screen Sharing: On", "Keep Screen Sharing off unless it is intentionally needed."),
        ]

    def _accessibility_checklist(self) -> list[FamilySafetyFinding]:
        items = [
            ("VoiceOver", "Recommended when screen reader support helps the user."),
            ("Zoom", "Recommended when magnification helps reading or focus."),
            ("Display scaling", "Recommended for users who benefit from larger interface elements."),
            ("Reduce motion", "Recommended for users sensitive to animation."),
            ("Reduce transparency", "Recommended for clarity and contrast."),
            ("Keyboard accessibility", "Recommended for users who navigate without a mouse."),
            ("Switch control", "Optional unless switch hardware is used."),
            ("Voice control", "Optional unless voice input improves access."),
            ("Assistive access", "Recommended for a simplified experience when appropriate."),
            ("Large cursor", "Recommended when pointer tracking is difficult."),
            ("Color filters", "Optional for color vision or contrast needs."),
        ]
        return [
            FamilySafetyFinding("Accessibility", title, "partially_configured", "Review with the user to decide whether this support is helpful.", recommendation)
            for title, recommendation in items
        ]

    def _parent_checklist(self, findings: list[FamilySafetyFinding]) -> list[FamilySafetyFinding]:
        titles = {
            "Family Sharing": "Family Sharing configured",
            "Screen Time": "Screen Time enabled",
            "Content Restrictions": "Content & Privacy Restrictions enabled",
            "App Store Restrictions": "App Store restrictions",
            "Safari Protection": "Safari restrictions",
            "Download Restrictions": "App installation controls",
            "Remote Access Review": "Remote login status",
            "Password Protection": "Password requirements",
            "Device Tracking": "Location sharing settings",
            "Backup Strategy": "",
        }
        by_title = {item.title: item for item in findings}
        checklist: list[FamilySafetyFinding] = []
        for label, source_title in titles.items():
            if source_title and source_title in by_title:
                source = by_title[source_title]
                checklist.append(FamilySafetyFinding("Parent Checklist", label, source.status, source.summary, source.recommendation, source.evidence))
            else:
                checklist.append(FamilySafetyFinding("Parent Checklist", label, "partially_configured", "Confirm backups are enabled and recoverable.", "Use Time Machine or another trusted backup strategy."))
        return checklist

    def _safe_browsing_status(self) -> list[FamilySafetyFinding]:
        safari = self.home / "Library" / "Preferences" / "com.apple.Safari.plist"
        prefs = self._read_plist(safari)
        return [
            self._safari_pref("Web Safety", "Safari Safe Browsing enabled", prefs, "WarnAboutFraudulentWebsites", True, "Turn on Safari Fraudulent Website Warning."),
            self._safari_pref("Web Safety", "Fraudulent Website Warning enabled", prefs, "WarnAboutFraudulentWebsites", True, "Turn on warnings for deceptive websites."),
            self._safari_pref("Web Safety", "Popup Blocking enabled", prefs, "WebKitJavaScriptCanOpenWindowsAutomatically", False, "Block pop-up windows in Safari."),
            self._verify("Downloads", "Download protections enabled", "Keep Gatekeeper enabled and only allow trusted downloads.", "Downloaded files and browsing activity are not inspected."),
            self._safari_pref("Privacy", "Cookie/privacy protections", prefs, "BlockStoragePolicy", None, "Review Safari Privacy settings and prevent cross-site tracking."),
        ]

    def _app_review(self) -> list[FamilySafetyFinding]:
        app_dirs = [Path("/Applications"), self.home / "Applications"]
        findings: list[FamilySafetyFinding] = []
        for app_dir in app_dirs:
            if not app_dir.exists():
                continue
            for app in sorted(app_dir.glob("*.app"))[:80]:
                status = "partially_configured"
                recommendation = "Mark trusted only after confirming the app is expected and age-appropriate."
                evidence = str(app)
                signature = self._run(["/usr/bin/codesign", "-dv", "--verbose=2", str(app)])
                if "Authority=" in signature:
                    status = "configured"
                    recommendation = "Signed app. Still review whether it belongs on this Mac."
                elif signature:
                    recommendation = "Unsigned or unusual app. Review before allowing a child or vulnerable user to use it."
                findings.append(FamilySafetyFinding("Application Review", app.name.removesuffix(".app"), status, "Installed app metadata only. No app contents or user documents are inspected.", recommendation, evidence))
        if not findings:
            findings.append(FamilySafetyFinding("Application Review", "Installed Apps", "partially_configured", "No applications were listed from standard app folders.", "Review installed apps in Finder or System Settings."))
        return findings[:60]

    def recommendations_for_profile(self, profile: str) -> dict[str, list[str]]:
        base = {
            "Young Child": ["Use a standard account, not an administrator account.", "Enable Screen Time, app limits, content restrictions, and web content filtering.", "Require approval before installing apps."],
            "Teen": ["Keep Screen Time and privacy settings transparent and age-appropriate.", "Review app permissions together.", "Teach phishing, social media, and location-sharing habits."],
            "Adult": ["Focus on privacy permissions, downloads, passwords, FileVault, firewall, and backups."],
            "Senior": ["Use simple settings, larger text, scam awareness, trusted contacts, and remote access review."],
            "Shared Family Computer": ["Use separate accounts for each person.", "Disable Guest access unless intentionally needed.", "Review downloads, web safety, and app inventory regularly."],
            "Special Needs User": ["Review accessibility supports with the user.", "Use simplified access where helpful.", "Reduce unexpected prompts, app clutter, and confusing permissions."],
            "School Device": ["Use managed accounts where appropriate.", "Review web filtering, app inventory, guest access, remote access, and classroom restrictions."],
            "Security Research Device": ["Complete the Security Research Device wizard and preserve authorization scope.", "Review FileVault, Secure Boot, SIP, software provenance, research network/DNS/VPN scope, evidence handling, backups, and disclosure contacts."],
            "Government Asset": ["Confirm the asset owner, authorization boundary, data classification, MDM enrollment, and currently approved macOS baseline.", "Review identity, audit retention, removable media, network scope, incident reporting, backup, and recovery with the system security officer."],
            "Doctor's Device": ["Protect availability and patient privacy without collecting patient content in MSAA.", "Review encryption, screen lock, approved clinical apps, session access, network segmentation, backups, downtime procedures, and incident contacts."],
            "Nurse's Workstation": ["Review shared-session controls, rapid lock, approved clinical applications, removable media, and emergency-access procedures.", "Coordinate changes with clinical operations so security changes do not disrupt patient care."],
            "Health Device": ["Confirm device purpose, vendor support, safety impact, network segmentation, update validation, inventory, and recovery procedures.", "Do not change a regulated or patient-care device without clinical owner and vendor approval."],
            "Lawyer's Device / Legal Asset": ["Review encryption, identity separation, secure communications, client/matter access, remote access, backups, retention, legal hold, and secure deletion.", "Do not place privileged or client-confidential content in MSAA notes or ordinary exports."],
        }
        return {"profile": [profile], "recommendations": base.get(profile, base["Shared Family Computer"])}

    def education_cards(self) -> list[SafetyEducationCard]:
        return [
            SafetyEducationCard("Cyberbullying", "Encourage saving evidence and talking to a trusted adult early.", "Agree on who the user can contact for help."),
            SafetyEducationCard("Phishing", "Unexpected messages asking for passwords, codes, or money should be treated carefully.", "Use bookmarks for important sites."),
            SafetyEducationCard("Scams", "Pressure, secrecy, and urgent payment requests are warning signs.", "Pause before paying or sharing information."),
            SafetyEducationCard("Unsafe Downloads", "Free tools, game cheats, and installers can include unwanted software.", "Download only from trusted sources."),
            SafetyEducationCard("Password Safety", "Unique passwords reduce damage if one site is compromised.", "Use a password manager and avoid sharing passwords."),
            SafetyEducationCard("Social Media Safety", "Privacy settings help, but posts can still be copied.", "Review followers, public posts, and direct messages."),
            SafetyEducationCard("Location Sharing", "Location sharing should be intentional and understood.", "Review Find My and app location access together."),
            SafetyEducationCard("Online Predators", "Adults should not ask children for secrets, private images, or hidden conversations.", "Create a no-punishment path for asking for help."),
            SafetyEducationCard("Fake Websites", "Look-alike domains and fake login pages are common.", "Type important addresses directly or use saved bookmarks."),
            SafetyEducationCard("AI Deepfake Awareness", "Images, voices, and videos can be faked.", "Verify surprising requests through another trusted channel."),
        ]

    def family_security_forecast(self) -> list[SafetyEducationCard]:
        return [
            SafetyEducationCard("Current scam reminder", "Gift card, bank, delivery, and account-lock messages are common scam formats.", "Verify through official apps or websites."),
            SafetyEducationCard("Apple safety reminder", "Keep macOS and Safari updated so built-in protections stay current.", "Review Software Update monthly."),
            SafetyEducationCard("Family check-in", "Safety settings work best when users understand why they exist.", "Schedule a calm review of apps, privacy, and screen time."),
        ]

    def privacy_notice(self) -> list[str]:
        return [
            "This center does not capture messages, screenshots, keystrokes, browsing history, private browsing data, microphone data, or camera data.",
            "Reports are generated locally and are not uploaded.",
            "The audit checks settings, local metadata, and user-confirmed safety items.",
        ]

    def _score(self, findings: list[FamilySafetyFinding]) -> FamilySafetyScore:
        values = {"configured": 1.0, "partially_configured": 0.5, "not_configured": 0.0}
        categories: dict[str, list[float]] = {category: [] for category in SAFETY_CATEGORIES}
        completed: list[str] = []
        recommended: list[str] = []
        for item in findings:
            if item.category in categories:
                categories[item.category].append(values.get(item.status, 0.5))
            if item.status == "configured":
                completed.append(item.title)
            else:
                recommended.append(item.recommendation)
        category_scores = {
            category: round((sum(values_) / len(values_)) * 100) if values_ else 50
            for category, values_ in categories.items()
        }
        score = round(sum(category_scores.values()) / len(category_scores))
        return FamilySafetyScore(score, category_scores, recommended[:12], completed[:12])

    def _caregiver_dashboard(self, score: FamilySafetyScore, findings: list[FamilySafetyFinding], app_review: list[FamilySafetyFinding]) -> dict[str, Any]:
        return {
            "safety_score": score.score,
            "recent_changes": "Run periodic reviews to compare changes over time.",
            "new_apps": [item.title for item in app_review[:10]],
            "remote_access_changes": [item.to_dict() for item in findings if "sharing" in item.title.lower() or "remote" in item.title.lower()],
            "account_changes": [item.to_dict() for item in findings if item.category == "Account Safety"],
            "safety_recommendations": score.recommended_improvements[:6],
        }

    def _exists_finding(self, category: str, title: str, path: Path, recommendation: str) -> FamilySafetyFinding:
        status = "partially_configured" if path.exists() else "not_configured"
        summary = "Related local settings were found." if path.exists() else "No local indicator was found."
        return FamilySafetyFinding(category, title, status, summary, recommendation, str(path))

    def _verify(self, category: str, title: str, recommendation: str, evidence: str) -> FamilySafetyFinding:
        return FamilySafetyFinding(category, title, "partially_configured", "Needs user confirmation in System Settings.", recommendation, evidence)

    def _guest_account_status(self) -> FamilySafetyFinding:
        output = self._run(["/usr/bin/dscl", ".", "-read", "/Users/Guest", "AuthenticationAuthority"])
        if output:
            return FamilySafetyFinding("Account Safety", "Guest account status", "partially_configured", "Guest account exists on this Mac.", "Disable Guest access unless it is intentionally needed.", output[:400])
        return FamilySafetyFinding("Account Safety", "Guest account status", "configured", "Guest account was not listed by the local directory query.", "No action needed unless guest access is intentionally required.")

    def _command_status(self, category: str, title: str, command: list[str], configured_text: str, recommendation: str) -> FamilySafetyFinding:
        output = self._run(command)
        status = "configured" if configured_text.lower() in output.lower() else "not_configured"
        summary = "Configured." if status == "configured" else "Not configured or could not be verified."
        return FamilySafetyFinding(category, title, status, summary, recommendation, output[:400])

    def _sharing_status(self, category: str, title: str, enabled_text: str, recommendation: str) -> FamilySafetyFinding:
        output = self._run(["/usr/sbin/systemsetup", "-getremotelogin"]) if "Remote Login" in enabled_text else self._run(["/usr/sbin/system_profiler", "SPConfigurationProfileDataType"])
        status = "not_configured" if enabled_text.lower() in output.lower() else "configured"
        summary = "Potentially enabled; review this setting." if status == "not_configured" else "No enabled indicator found."
        return FamilySafetyFinding(category, title, status, summary, recommendation, output[:400])

    def _safari_pref(self, category: str, title: str, prefs: dict[str, Any], key: str, expected: Any, recommendation: str) -> FamilySafetyFinding:
        if not prefs:
            return FamilySafetyFinding(category, title, "partially_configured", "Safari preferences were not available for local review.", recommendation)
        value = prefs.get(key)
        if expected is None:
            status = "configured" if value is not None else "partially_configured"
        else:
            status = "configured" if value == expected else "not_configured"
        return FamilySafetyFinding(category, title, status, f"Preference {key}: {value}", recommendation)

    def _read_plist(self, path: Path) -> dict[str, Any]:
        try:
            with path.open("rb") as handle:
                payload = plistlib.load(handle)
            return payload if isinstance(payload, dict) else {}
        except (OSError, plistlib.InvalidFileException):
            return {}

    def _run(self, command: list[str]) -> str:
        try:
            result = subprocess.run(command, capture_output=True, text=True, timeout=8, check=False)
        except (OSError, subprocess.SubprocessError):
            return ""
        return "\n".join(part for part in [result.stdout.strip(), result.stderr.strip()] if part)


def default_family_safety_report_path(base_dir: Path | None = None, suffix: str = "html", now: datetime | None = None) -> Path:
    reports_dir = (base_dir / "reports") if base_dir is not None else Path.home() / "Library" / "Application Support" / "MacAuditAgent" / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)
    timestamp = (now or datetime.now()).strftime("%Y%m%d_%H%M%S")
    return reports_dir / f"family_safety_report_{timestamp}.{suffix}"


def export_family_safety_json(report: FamilySafetyReport, output_path: Path | None = None) -> Path:
    output_path = output_path or default_family_safety_report_path(suffix="json")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(report.to_dict(), indent=2), encoding="utf-8")
    return output_path


def export_family_safety_html(report: FamilySafetyReport, output_path: Path | None = None) -> Path:
    output_path = output_path or default_family_safety_report_path(suffix="html")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    payload = report.to_dict()
    score = payload["score"]
    findings = payload["findings"]
    missing = [item for item in findings if item.get("status") != "configured"]
    completed = [item for item in findings if item.get("status") == "configured"]
    education = payload["education_cards"]
    privacy = payload["privacy_notice"]
    categories = payload.get("category_definitions", [])
    category_scores = payload.get("category_scores", {})
    limitations = payload.get("limitations", [])

    def rows(items: list[dict[str, Any]], fields: list[str]) -> str:
        rendered = []
        for item in items:
            rendered.append("<tr>" + "".join(f"<td>{html.escape(str(item.get(field, '')))}</td>" for field in fields) + "</tr>")
        return "".join(rendered)

    output_path.write_text(
        f"""<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8">
  <title>Family Safety Report</title>
  <style>
    body {{ font-family: -apple-system, BlinkMacSystemFont, sans-serif; margin: 32px; color: #172033; }}
    h1, h2 {{ color: #0f172a; }}
    .score {{ font-size: 32px; font-weight: 700; margin: 12px 0; }}
    .notice {{ background: #eef6ff; border: 1px solid #bfdbfe; padding: 12px; border-radius: 8px; }}
    table {{ width: 100%; border-collapse: collapse; margin: 12px 0 24px; }}
    th, td {{ border: 1px solid #d7dee8; padding: 8px; text-align: left; vertical-align: top; }}
    th {{ background: #f1f5f9; }}
  </style>
</head>
<body>
  <h1>Family Safety Report</h1>
  <p>Generated locally: {html.escape(str(payload.get("generated_at", "")))}</p>
  <div class="score">Safety Score: {html.escape(str(score.get("score", "--")))}/100</div>
  <div class="notice">
    <strong>Privacy-first report.</strong>
    <ul>{''.join(f'<li>{html.escape(str(item))}</li>' for item in privacy)}</ul>
  </div>
  <h2>Recommended Improvements</h2>
  <ul>{''.join(f'<li>{html.escape(str(item))}</li>' for item in score.get("recommended_improvements", []))}</ul>
  <h2>Family Control Center Categories</h2>
  <p>These sections are plain-language local review guidance. NIST references are NIST-aligned mappings for review support only.</p>
  <table><tr><th>Category</th><th>Score</th><th>Status</th><th>Description</th><th>Why It Matters</th><th>Recommended Changes</th></tr>{rows([
      {
          "title": item.get("title", ""),
          "score": category_scores.get(item.get("category_id", ""), ""),
          "current_status": item.get("current_status", ""),
          "short_description": item.get("short_description", ""),
          "why_it_matters": item.get("why_it_matters", ""),
          "what_the_user_can_change": "; ".join(item.get("what_the_user_can_change", [])),
      }
      for item in categories
  ], ["title", "score", "current_status", "short_description", "why_it_matters", "what_the_user_can_change"])}</table>
  <h2>Government / NIST Lockdown Profile</h2>
  <p>Government Lockdown Score: {html.escape(str(payload.get("government_lockdown_score", 0)))}/100. This is hardening guidance mapped to NIST references and does not claim compliance or certification.</p>
  <h2>Lockdown Mode Plus</h2>
  <p>Lockdown Plus Score: {html.escape(str(payload.get("lockdown_plus_score", 0)))}/100. Status: {html.escape(str(payload.get("lockdown_plus_status", "")))}. Lockdown Mode Plus complements Apple Lockdown Mode guidance but does not replace it.</p>
  <h2>Completed Actions</h2>
  <ul>{''.join(f'<li>{html.escape(str(item.get("title", "")))}</li>' for item in completed)}</ul>
  <h2>Missing or Needs Review</h2>
  <table><tr><th>Category</th><th>Check</th><th>Status</th><th>Next Step</th></tr>{rows(missing, ["category", "title", "status", "recommendation"])}</table>
  <h2>Full Safety Audit</h2>
  <table><tr><th>Category</th><th>Check</th><th>Status</th><th>Guidance</th></tr>{rows(findings, ["category", "title", "status", "recommendation"])}</table>
  <h2>Online Safety Guidance</h2>
  <table><tr><th>Topic</th><th>Guidance</th><th>Action</th></tr>{rows(education, ["topic", "guidance", "action"])}</table>
  <h2>Limitations</h2>
  <ul>{''.join(f'<li>{html.escape(str(item))}</li>' for item in limitations)}</ul>
</body>
</html>
""",
        encoding="utf-8",
    )
    return output_path


def export_family_safety_word(report: FamilySafetyReport, output_path: Path | None = None) -> Path:
    try:
        from docx import Document
    except Exception:
        from mac_audit_agent.professional_report import structured_payload_report

        output_path = output_path or default_family_safety_report_path(suffix="docx")
        return structured_payload_report(
            output_path,
            title="Family Safety Report",
            payload=report.to_dict(),
            qualification="Static macro-free Word report; findings and recommendations require authorized human review.",
        )
    output_path = output_path or default_family_safety_report_path(suffix="docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    payload = report.to_dict()
    doc = Document()
    doc.add_heading("Family Safety Report", 0)
    doc.add_paragraph(f"Generated locally: {payload.get('generated_at', '')}")
    doc.add_paragraph(f"Family Safety Score: {payload.get('score', {}).get('score', '--')}/100")
    doc.add_heading("Privacy Notice", 1)
    for item in payload.get("privacy_notice", []):
        doc.add_paragraph(str(item), style="List Bullet")
    doc.add_heading("Family Control Center Categories", 1)
    for category in payload.get("category_definitions", []):
        doc.add_heading(str(category.get("title", "")), 2)
        doc.add_paragraph(str(category.get("short_description", "")))
        doc.add_paragraph(f"Status: {category.get('current_status', '')}")
        doc.add_paragraph(f"Why this matters: {category.get('why_it_matters', '')}")
        doc.add_paragraph("Recommended changes:")
        for item in category.get("what_the_user_can_change", []):
            doc.add_paragraph(str(item), style="List Bullet")
    doc.add_heading("Government / NIST Lockdown Profile", 1)
    doc.add_paragraph(
        f"Government Lockdown Score: {payload.get('government_lockdown_score', 0)}/100. "
        "This is NIST-aligned hardening guidance for review support only."
    )
    doc.add_heading("Lockdown Mode Plus", 1)
    doc.add_paragraph(
        f"Lockdown Plus Score: {payload.get('lockdown_plus_score', 0)}/100. "
        f"Status: {payload.get('lockdown_plus_status', '')}. "
        "This complements Apple Lockdown Mode guidance but does not replace it."
    )
    doc.add_heading("Limitations", 1)
    for item in payload.get("limitations", []):
        doc.add_paragraph(str(item), style="List Bullet")
    with NamedTemporaryFile(prefix=output_path.stem, suffix=".tmp.docx", dir=output_path.parent, delete=False) as handle:
        temp_path = Path(handle.name)
    try:
        doc.save(temp_path)
        temp_path.replace(output_path)
    except Exception:
        temp_path.unlink(missing_ok=True)
        raise
    return output_path


def export_family_safety_excel(report: FamilySafetyReport, output_path: Path | None = None) -> Path:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
    except Exception:
        from mac_audit_agent.professional_report import structured_payload_report

        output_path = output_path or default_family_safety_report_path(suffix="xlsx")
        return structured_payload_report(
            output_path,
            title="Family Safety Report",
            payload=report.to_dict(),
            qualification="Static formula-free Excel workbook; findings and recommendations require authorized human review.",
        )
    output_path = output_path or default_family_safety_report_path(suffix="xlsx")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    payload = report.to_dict()
    wb = Workbook()
    ws = wb.active
    ws.title = "Family Summary"
    ws.append(["Field", "Value"])
    ws.append(["Family Safety Score", payload.get("score", {}).get("score", "")])
    ws.append(["Government Lockdown Score", payload.get("government_lockdown_score", 0)])
    ws.append(["Lockdown Plus Score", payload.get("lockdown_plus_score", 0)])
    ws.append(["Lockdown Plus Status", payload.get("lockdown_plus_status", "")])
    categories_ws = wb.create_sheet("Categories")
    columns = ["category_id", "title", "current_status", "score", "short_description", "why_it_matters", "recommended_changes", "nist_mappings"]
    categories_ws.append(columns)
    scores = payload.get("category_scores", {})
    for category in payload.get("category_definitions", []):
        categories_ws.append(
            [
                category.get("category_id", ""),
                category.get("title", ""),
                category.get("current_status", ""),
                scores.get(category.get("category_id", ""), ""),
                category.get("short_description", ""),
                category.get("why_it_matters", ""),
                "; ".join(category.get("what_the_user_can_change", [])),
                "; ".join(category.get("nist_mappings", [])),
            ]
        )
    checklist_ws = wb.create_sheet("Checklist")
    checklist_ws.append(["Category", "Checklist Item"])
    for category in payload.get("category_definitions", []):
        for item in category.get("checklist_items", []):
            checklist_ws.append([category.get("title", ""), item])
    limitations_ws = wb.create_sheet("Limitations")
    limitations_ws.append(["Limitation"])
    for item in payload.get("limitations", []):
        limitations_ws.append([item])
    header_fill = PatternFill("solid", fgColor="1F4E78")
    for sheet in wb.worksheets:
        for cell in sheet[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = header_fill
        for row in sheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")
        for column_cells in sheet.columns:
            width = min(60, max(12, max(len(str(cell.value or "")) for cell in column_cells[:100]) + 2))
            sheet.column_dimensions[column_cells[0].column_letter].width = width
    with NamedTemporaryFile(prefix=output_path.stem, suffix=".tmp.xlsx", dir=output_path.parent, delete=False) as handle:
        temp_path = Path(handle.name)
    try:
        wb.save(temp_path)
        temp_path.replace(output_path)
    except Exception:
        temp_path.unlink(missing_ok=True)
        raise
    return output_path


from mac_audit_agent.family_safety.apply_engine import (  # noqa: E402
    FamilySafetyConfigSnapshot,
    apply_family_safety_recommendation,
    latest_family_safety_snapshot,
    restore_family_safety_snapshot,
)
from mac_audit_agent.family_safety.config_change import FamilySafetyConfigChange  # noqa: E402
from mac_audit_agent.family_safety.profiles import FamilySafetyProfile, canonical_family_safety_profiles  # noqa: E402
from mac_audit_agent.family_safety.recommendation_engine import FamilySafetyRecommendation, FamilySafetyRecommendationEngine  # noqa: E402
from mac_audit_agent.family_safety.reporting import (  # noqa: E402
    export_family_safety_configuration_excel,
    export_family_safety_configuration_html,
    export_family_safety_configuration_json,
    export_family_safety_configuration_markdown,
    export_family_safety_configuration_word,
)
from mac_audit_agent.family_safety.wizard_questions import FamilySafetyQuestion, canonical_family_safety_questions  # noqa: E402

from __future__ import annotations

import html
import json
from dataclasses import asdict, dataclass
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any
from uuid import uuid4

from mac_audit_agent.models import BackgroundMonitorEvent, utc_now_iso


STANDARD_SUGGESTIONS = (
    ("NIST CSF Govern", "Confirm scope, authorization, stakeholders, risk decisions, and evidence ownership."),
    ("NIST CSF Identify", "Inventory assets, software, identities, dependencies, and material exposure."),
    ("NIST CSF Protect", "Review access control, data protection, platform hardening, and secure configuration."),
    ("NIST CSF Detect", "Validate telemetry coverage, alert paths, baselines, and detection evidence."),
    ("NIST CSF Respond", "Review triage, communications, containment authorization, and evidence preservation."),
    ("NIST CSF Recover", "Validate rollback, restoration, recovery evidence, and lessons learned."),
    ("NIST SP 800-171 / CMMC", "Map observed evidence to the configured requirement version; record gaps without claiming certification."),
    ("DoD mission assurance", "Confirm mission impact, authorization boundaries, recovery needs, and accountable approval."),
)


@dataclass(frozen=True)
class TimesheetEntry:
    entry_id: str
    contractor_name: str
    engagement_name: str
    started_at: str
    ended_at: str
    duration_seconds: int
    work_notes: str
    goals: str
    completed: str
    struggles: str
    tools_used: str
    standards_focus: str
    created_at: str
    updated_at: str

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


class ConsultantTimesheetRepository:
    def __init__(self, audit_database) -> None:
        self.db = audit_database
        self.db.conn.executescript(
            """
            CREATE TABLE IF NOT EXISTS consultant_timesheet_entries (
                entry_id TEXT PRIMARY KEY,
                contractor_name TEXT NOT NULL,
                engagement_name TEXT NOT NULL,
                started_at TEXT NOT NULL,
                ended_at TEXT NOT NULL DEFAULT '',
                duration_seconds INTEGER NOT NULL DEFAULT 0,
                work_notes TEXT NOT NULL DEFAULT '',
                goals TEXT NOT NULL DEFAULT '',
                completed TEXT NOT NULL DEFAULT '',
                struggles TEXT NOT NULL DEFAULT '',
                tools_used TEXT NOT NULL DEFAULT '',
                standards_focus TEXT NOT NULL DEFAULT '',
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            );
            CREATE INDEX IF NOT EXISTS idx_consultant_timesheet_started ON consultant_timesheet_entries(started_at);
            """
        )
        self.db.conn.commit()

    @staticmethod
    def _clean(value: str, limit: int = 20000) -> str:
        return str(value or "").replace("\x00", "").strip()[:limit]

    def active(self) -> TimesheetEntry | None:
        row = self.db.conn.execute("SELECT * FROM consultant_timesheet_entries WHERE ended_at = '' ORDER BY started_at DESC LIMIT 1").fetchone()
        return self._from_row(row) if row else None

    def start(self, contractor_name: str, engagement_name: str, *, timestamp: str | None = None) -> TimesheetEntry:
        if self.active() is not None:
            raise ValueError("An assessment is already clocked in.")
        contractor = self._clean(contractor_name, 300)
        engagement = self._clean(engagement_name, 300)
        if not contractor or not engagement:
            raise ValueError("Contractor/work name and engagement name are required.")
        now = timestamp or utc_now_iso()
        entry_id = f"timesheet-{uuid4().hex}"
        self.db.conn.execute(
            "INSERT INTO consultant_timesheet_entries(entry_id,contractor_name,engagement_name,started_at,created_at,updated_at) VALUES(?,?,?,?,?,?)",
            (entry_id, contractor, engagement, now, now, now),
        )
        self.db.conn.commit()
        self._event("consultant_timesheet_started", entry_id, now, f"Billable assessment timer started for engagement {engagement}.", {"contractor_name": contractor, "engagement_name": engagement})
        return self.get(entry_id)

    def save_notes(self, entry_id: str, *, work_notes: str, goals: str, completed: str, struggles: str, tools_used: str, standards_focus: str) -> TimesheetEntry:
        now = utc_now_iso()
        values = tuple(self._clean(value) for value in (work_notes, goals, completed, struggles, tools_used, standards_focus))
        cursor = self.db.conn.execute(
            "UPDATE consultant_timesheet_entries SET work_notes=?,goals=?,completed=?,struggles=?,tools_used=?,standards_focus=?,updated_at=? WHERE entry_id=?",
            (*values, now, entry_id),
        )
        if cursor.rowcount != 1:
            raise ValueError("Timesheet entry was not found.")
        self.db.conn.commit()
        self._event("consultant_timesheet_notes_saved", entry_id, now, "Timesheet work notes and goals were saved.", {"standards_focus": values[-1]})
        return self.get(entry_id)

    def stop(self, entry_id: str, *, timestamp: str | None = None) -> TimesheetEntry:
        entry = self.get(entry_id)
        if entry.ended_at:
            return entry
        ended = timestamp or utc_now_iso()
        start_dt = datetime.fromisoformat(entry.started_at.replace("Z", "+00:00"))
        end_dt = datetime.fromisoformat(ended.replace("Z", "+00:00"))
        duration = max(0, round((end_dt - start_dt).total_seconds()))
        self.db.conn.execute("UPDATE consultant_timesheet_entries SET ended_at=?,duration_seconds=?,updated_at=? WHERE entry_id=? AND ended_at=''", (ended, duration, ended, entry_id))
        self.db.conn.commit()
        self._event("consultant_timesheet_stopped", entry_id, ended, f"Billable assessment timer stopped after {duration} seconds.", {"duration_seconds": duration})
        return self.get(entry_id)

    def get(self, entry_id: str) -> TimesheetEntry:
        row = self.db.conn.execute("SELECT * FROM consultant_timesheet_entries WHERE entry_id=?", (entry_id,)).fetchone()
        if not row:
            raise ValueError("Timesheet entry was not found.")
        return self._from_row(row)

    def list_entries(self, *, since: str = "", until: str = "") -> list[TimesheetEntry]:
        query = "SELECT * FROM consultant_timesheet_entries WHERE 1=1"
        parameters: list[str] = []
        if since:
            query += " AND started_at >= ?"; parameters.append(since)
        if until:
            query += " AND started_at < ?"; parameters.append(until)
        query += " ORDER BY started_at DESC"
        return [self._from_row(row) for row in self.db.conn.execute(query, parameters).fetchall()]

    def record_export(self, *, period: str, format_name: str, entry_count: int, output_name: str) -> None:
        now = utc_now_iso()
        self._event("consultant_timesheet_exported", "timesheet-export", now, f"Consultant timesheet export created for {period}.", {"period": period, "format": format_name, "entry_count": entry_count, "output_name": Path(output_name).name})

    def _event(self, event_type: str, entry_id: str, timestamp: str, evidence: str, metadata: dict[str, Any]) -> None:
        payload = {"schema_version": "1.0", "entry_id": entry_id, **metadata, "qualification": "User-initiated timekeeping record; not proof of continuous activity or contractual acceptance."}
        self.db.record_background_monitor_event(BackgroundMonitorEvent(event_id=f"{event_type}-{uuid4().hex}", timestamp=timestamp, event_type=event_type, severity="info", source="consultant_timesheet", evidence=evidence, confidence="high", metadata_json=json.dumps(payload, sort_keys=True), notification_decision="log_only"), dedupe_window_seconds=0)

    @staticmethod
    def _from_row(row) -> TimesheetEntry:
        return TimesheetEntry(**{key: row[key] for key in TimesheetEntry.__dataclass_fields__})


def range_bounds(period: str, now: datetime | None = None) -> tuple[str, str]:
    current = now or datetime.now(timezone.utc)
    if current.tzinfo is None:
        current = current.replace(tzinfo=timezone.utc)
    if period == "Daily": start = current.replace(hour=0, minute=0, second=0, microsecond=0)
    elif period == "Weekly": start = (current - timedelta(days=current.weekday())).replace(hour=0, minute=0, second=0, microsecond=0)
    elif period == "Monthly": start = current.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    else: return "", ""
    if period == "Daily": end = start + timedelta(days=1)
    elif period == "Weekly": end = start + timedelta(days=7)
    else: end = (start.replace(day=28) + timedelta(days=4)).replace(day=1)
    return start.isoformat(), end.isoformat()


def export_timesheet(entries: list[TimesheetEntry], output: Path, *, title: str = "Consultant Timesheet") -> Path:
    suffix = output.suffix.lower()
    rows = [entry.to_dict() for entry in entries]
    total = sum(entry.duration_seconds for entry in entries)
    output.parent.mkdir(parents=True, exist_ok=True)
    if suffix == ".txt":
        lines = [title, f"Total hours: {total / 3600:.2f}", ""]
        for entry in entries: lines.extend([f"{entry.started_at} — {entry.engagement_name} — {entry.duration_seconds / 3600:.2f} h", f"Contractor: {entry.contractor_name}", f"Notes: {entry.work_notes}", f"Goals: {entry.goals}", f"Completed: {entry.completed}", f"Struggles: {entry.struggles}", f"Tools: {entry.tools_used}", f"Standards focus: {entry.standards_focus}", ""])
        output.write_text("\n".join(lines), encoding="utf-8")
    elif suffix == ".html":
        body = "".join("<tr>" + "".join(f"<td>{html.escape(str(row[key]))}</td>" for key in ("started_at","ended_at","duration_seconds","contractor_name","engagement_name","work_notes","goals","completed","struggles","tools_used","standards_focus")) + "</tr>" for row in rows)
        output.write_text(f"<!doctype html><meta charset='utf-8'><title>{html.escape(title)}</title><h1>{html.escape(title)}</h1><p>Total hours: {total/3600:.2f}</p><table border='1'><thead><tr><th>Start</th><th>End</th><th>Seconds</th><th>Contractor</th><th>Engagement</th><th>Notes</th><th>Goals</th><th>Completed</th><th>Struggles</th><th>Tools</th><th>Standards</th></tr></thead><tbody>{body}</tbody></table>", encoding="utf-8")
    elif suffix == ".xlsx":
        from openpyxl import Workbook
        workbook = Workbook(); sheet = workbook.active; sheet.title = "Timesheet"; headers = list(TimesheetEntry.__dataclass_fields__); sheet.append(headers)
        for row in rows: sheet.append([row[key] for key in headers])
        sheet.freeze_panes = "A2"; workbook.save(output)
    elif suffix == ".docx":
        from docx import Document
        document = Document(); document.add_heading(title, 0); document.add_paragraph(f"Total hours: {total/3600:.2f}")
        for entry in entries:
            document.add_heading(f"{entry.engagement_name} — {entry.started_at}", level=1)
            for label, value in (("Contractor",entry.contractor_name),("End",entry.ended_at or "Active"),("Hours",f"{entry.duration_seconds/3600:.2f}"),("Work notes",entry.work_notes),("Goals",entry.goals),("Completed",entry.completed),("Struggles",entry.struggles),("Tools",entry.tools_used),("Standards focus",entry.standards_focus)): document.add_paragraph(f"{label}: {value}")
        document.save(output)
    elif suffix == ".pdf":
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        pdf = canvas.Canvas(str(output), pagesize=letter); width, height = letter; y = height - 54; pdf.setFont("Helvetica-Bold", 16); pdf.drawString(54, y, title); y -= 24; pdf.setFont("Helvetica", 10); pdf.drawString(54, y, f"Total hours: {total/3600:.2f}")
        for entry in entries:
            for line in (f"{entry.started_at} | {entry.engagement_name} | {entry.duration_seconds/3600:.2f} h", f"Contractor: {entry.contractor_name}", f"Notes: {entry.work_notes}", f"Goals: {entry.goals}", f"Completed: {entry.completed}", f"Struggles: {entry.struggles}", f"Tools: {entry.tools_used}", f"Standards: {entry.standards_focus}"):
                if y < 54: pdf.showPage(); y = height - 54; pdf.setFont("Helvetica", 10)
                pdf.drawString(54, y, line[:110]); y -= 14
            y -= 8
        pdf.save()
    else:
        raise ValueError("Supported formats are .xlsx, .docx, .pdf, .txt, and .html")
    return output
